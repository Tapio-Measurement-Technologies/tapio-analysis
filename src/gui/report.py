from PyQt6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QLabel, QComboBox, QSizePolicy,
                             QPushButton, QScrollArea, QFrame, QLineEdit, QFileDialog, QMenuBar, QTextEdit, QMessageBox)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QStandardItemModel, QStandardItem, QAction
from utils.data_loader import DataMixin
from controllers import *
from utils.windows import *
import settings
import datetime
from docx import Document
from docx.shared import Mm, Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn



import numpy as np
import json
import os
from customizations import apply_plot_customizations
from utils.report import get_text_width, set_paragraph_spacing
import traceback
import importlib.util
import matplotlib.pyplot as plt


class Editor(QTextEdit):
    def __init__(self):
        super().__init__()
        self.textChanged.connect(self.autoResize)

    def autoResize(self):
        self.document().setTextWidth(self.viewport().width())
        margins = self.contentsMargins()
        height = int(self.document().size().height() +
                     margins.top() + margins.bottom())
        self.setFixedHeight(height)

    def resizeEvent(self, event):
        self.autoResize()


class ReportWindow(QWidget, DataMixin):
    def __init__(self, main_window, window_type="MD"):
        super().__init__()
        self.dataMixin = DataMixin.getInstance()
        self.setWindowTitle(f"Generate {window_type} Report ({
                            self.dataMixin.measurement_label})")
        self.setGeometry(100, 100, 700, 800)
        self.main_window = main_window
        self.window_type = window_type
        self.report_title = f"{window_type} Report"
        self.report_subtitle = f"Mill PM 1"
        self.section_widgets = []
        self.header_image_path = settings.REPORT_HEADER_IMAGE_PATH

        # Main layout
        self.main_layout = QVBoxLayout()
        self.main_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.setLayout(self.main_layout)

        # Menu
        self.initMenuBar(self.main_layout)

        # Title input
        self.title_layout = QVBoxLayout()
        self.title_label = QLabel("Report Title:")
        self.title_input = QLineEdit(self.report_title)
        self.title_input.textEdited.connect(self.update_report_title)
        self.title_layout.addWidget(self.title_label)
        self.title_layout.addWidget(self.title_input)
        self.main_layout.addLayout(self.title_layout)

        # Title input
        self.subtitle_layout = QVBoxLayout()
        self.subtitle_label = QLabel("Report subtitle:")
        self.subtitle_input = QLineEdit(self.report_subtitle)
        self.subtitle_input.textEdited.connect(self.update_report_subtitle)
        self.subtitle_layout.addWidget(self.subtitle_label)
        self.subtitle_layout.addWidget(self.subtitle_input)
        self.main_layout.addLayout(self.subtitle_layout)

        # Image input
        self.header_image_layout = QHBoxLayout()
        self.header_image_label = QLabel("Header Image:")
        self.header_image_path_input = QLineEdit(self.header_image_path)
        self.header_image_path_button = QPushButton("Choose Image")
        self.header_image_path_button.clicked.connect(self.choose_image)
        self.header_image_layout.addWidget(self.header_image_label)
        self.header_image_layout.addWidget(self.header_image_path_input)
        self.header_image_layout.addWidget(self.header_image_path_button)
        self.main_layout.addLayout(self.header_image_layout)

        # Additional information input
        self.additional_info_layout = QVBoxLayout()
        self.additional_info_label = QLabel("Additional Information:")
        self.additional_info_input = Editor()
        self.additional_info_input.setText(
            settings.REPORT_ADDITIONAL_INFO_DEFAULT)
        self.additional_info_input.setSizePolicy(
            QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Minimum)
        self.additional_info_layout.addWidget(self.additional_info_label)
        self.additional_info_layout.addWidget(self.additional_info_input)
        self.main_layout.addLayout(self.additional_info_layout)

        # Scroll area
        self.scroll_area = QScrollArea(self)
        self.scroll_area.setWidgetResizable(True)
        self.main_layout.addWidget(self.scroll_area)

        # Widget inside scroll area
        self.scroll_widget = QWidget()
        self.scroll_area.setWidget(self.scroll_widget)

        self.scroll_layout = QVBoxLayout(self.scroll_widget)
        self.scroll_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        # Heading
        heading = QLabel("Sections")
        heading.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.scroll_layout.addWidget(heading)

        # Container for sections
        self.sections_container = QVBoxLayout()
        self.sections_container.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.scroll_layout.addLayout(self.sections_container)

        # Button for adding sections
        self.add_section_button = QPushButton("+ Add section")
        self.add_section_button.clicked.connect(lambda: self.add_section())
        self.scroll_layout.addWidget(self.add_section_button)

        # Generate Report Button
        self.generate_report_button = QPushButton("Generate Report")
        self.generate_report_button.clicked.connect(self.generate_report)
        self.main_layout.addWidget(self.generate_report_button)

    def initMenuBar(self, layout):
        menuBar = QMenuBar()
        layout.setMenuBar(menuBar)

        fileMenu = menuBar.addMenu('File')
        loadFileAction = QAction('Load report template', self)
        loadFileAction.setShortcut("Ctrl+O")
        loadFileAction.setStatusTip("Open file")
        loadFileAction.triggered.connect(self.load_from_python)
        fileMenu.addAction(loadFileAction)

    def add_section(self, section_name="Section"):
        section_widget = ReportSectionWidget(
            self.main_window, section_name, self.window_type)
        self.section_widgets.append(section_widget)
        section_widget.destroyed.connect(
            lambda: self.section_widgets.remove(section_widget))
        self.sections_container.addWidget(section_widget)
        return section_widget

    def update_report_title(self, title):
        self.report_title = title
        self.title_input.setText(self.report_title)

    def update_report_subtitle(self, subtitle):
        self.report_subtitle = subtitle
        self.subtitle_input.setText(self.report_subtitle)

    def choose_image(self):
        dialog = QFileDialog(self)
        options = QFileDialog.options(dialog)
        file_name, _ = QFileDialog.getOpenFileName(
            self, "Choose Image", "", "Image Files (*.png *.jpg *.bmp *.gif)", options=options)
        if file_name:
            self.header_image_path = file_name
            self.header_image_path_input.setText(file_name)

    def generate_report(self):
        doc = Document()
        margin = Cm(1)
        for section in doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        style = doc.styles['Normal']
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        # Explicitly define font in Word
        r = style._element
        r_rPr = r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "DejaVu Sans")
        rFonts.set(qn("w:hAnsi"), "DejaVu Sans")
        rFonts.set(qn("w:eastAsia"), "DejaVu Sans")
        rFonts.set(qn("w:cs"), "DejaVu Sans")
        r_rPr.append(rFonts)

        # Header section
        header = doc.sections[0].header
        header_table = header.add_table(
            rows=1, cols=3, width=Mm(get_text_width(doc) - 60))
        header_table.autofit = False

        col1_width = Mm(65)
        # Dynamic width for center text
        col2_width = Mm(60)
        col3_width = Mm(60)

        header_table.columns[0].width = col1_width
        header_table.columns[1].width = col2_width
        header_table.columns[2].width = col3_width

        # Access cells
        cell1 = header_table.cell(0, 0)
        cell2 = header_table.cell(0, 1)
        cell3 = header_table.cell(0, 2)

        # Add left image
        if self.header_image_path:
            run = cell1.paragraphs[0].add_run()
            run.add_picture(self.header_image_path, width=Mm(50))


        p = cell2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Clear existing text
        p.clear()

        # Add bold title
        run_title = p.add_run(self.report_title)
        run_title.bold = True  # Make only the title bold

        # Add subtitle and timestamp (not bold)
        run_subtitle = p.add_run(f"\n{self.report_subtitle}\n\nGenerated {str(datetime.datetime.now())}")
        run_subtitle.bold = False  # Ensure this part is not bold


        # cell1.paragraphs[0].text = (
        #     f"Generated {str(datetime.datetime.now())}")


        # Add right image (if applicable)
        roll_image_path = settings.MD_REPORT_HEADER_IMAGE_PATH if self.window_type == "MD" else settings.CD_REPORT_HEADER_IMAGE_PATH
        if roll_image_path:
            run = cell3.paragraphs[0].add_run()
            run.add_picture(roll_image_path, width=Mm(50))
            cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


        for section in self.section_widgets:
            paragraph = doc.add_heading(section.section_name)

            run = paragraph.runs[0]
            run.font.color.rgb = None  # Removes the blue color (defaults to black)
            run.font.size = Pt(12)  # Adjust size (default heading sizes are often too big)
            run.font.name = "Calibri"  # Ensure a consistent font


            set_paragraph_spacing(paragraph, 0, 6)
            for analysis in section.analysis_widgets:
                apply_plot_customizations(analysis)

                analysis_title = analysis.analysis_title or f"{analysis.analysis_name} {
                    analysis.get_channel_text()}"
                paragraph = doc.add_heading(analysis_title, 2)
                run = paragraph.runs[0]
                run.font.color.rgb = None  # Set to black
                run.font.size = Pt(11)  # Slightly smaller for subsection headings
                run.font.name = "Calibri"

                # paragraph = doc.add_heading(f"{analysis.analysis_name} {
                #                             analysis.get_channel_text()}", 2)
                set_paragraph_spacing(paragraph)

                # layout_mode = ["stats-right", "stats-below", "stats-above"]
                # Default to stats-right
                layout_mode = analysis.report_layout or "stats-right"

                # Define column widths for side-by-side layout
                img_col_width = Mm(get_text_width(doc) * (3/5))
                stats_col_width = Mm(get_text_width(doc) * (2/5))

                if layout_mode == "stats-right":
                    # Side-by-side layout (image left, stats right)
                    table = doc.add_table(rows=1, cols=2)
                    table.autofit = True

                    col1 = table.columns[0]
                    col1.width = img_col_width
                    cell1 = table.cell(0, 0)
                    cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    paragraph = cell1.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(
                        analysis.controller.getPlotImage(), width=Cm(col1.width.cm - 0.5))

                    col2 = table.columns[1]
                    col2.width = stats_col_width
                    cell2 = table.cell(0, 1)
                    cell2.width = stats_col_width
                    cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    img_cell = table.cell(0, 0)
                    stats_cell = table.cell(0, 1)

                elif layout_mode in ["stats-below", "stats-above"]:
                    # Image and stats stacked vertically
                    table = doc.add_table(rows=2, cols=1)
                    table.autofit = True

                    if layout_mode == "stats-above":
                        stats_cell = table.cell(0, 0)
                        img_cell = table.cell(1, 0)
                    else:  # stats-below
                        img_cell = table.cell(0, 0)
                        stats_cell = table.cell(1, 0)

                    # Add image to the image cell
                    run = img_cell.paragraphs[0].add_run()
                    run.add_picture(analysis.controller.getPlotImage(
                    ), width=Mm(get_text_width(doc) - 5))

                    stats_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                # Retrieve statistics data
                data = analysis.controller.getStatsTableData()
                if not data:
                    pass  # No data, skip stats table
                else:
                    shape = np.shape(data)
                    rows, cols = shape if len(shape) == 2 else (
                        shape[0], 1) if len(shape) == 1 else (1, 1)

                    # Insert stats table
                    stats_table = stats_cell.add_table(rows, cols)
                    delete_paragraph(stats_cell.paragraphs[0])

                    for row_idx, row_data in enumerate(data):
                        row = stats_table.rows[row_idx]
                        for col_idx, cell_data in enumerate(row_data):
                            cell = row.cells[col_idx]
                            cell.paragraphs[0].text = cell_data
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            cell.paragraphs[0].style.font.size = Pt(8)
                            cell.paragraphs[0].style.font.name = "Nimbus Mono PS"

            # Add a page break after each section
            # TODO: this causes a problem where empty pages may appear

            # Only add a break if the last paragraph is not empty
            if doc.paragraphs[-1].text.strip():
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Open save file dialog
        dialog = QFileDialog()
        options = QFileDialog.options(dialog)
        fileName, _ = QFileDialog.getSaveFileName(
            self, "Word Document", "", "DOCX Files (*.docx)", options=options)
        if fileName:
            if not fileName.endswith('.docx'):
                fileName += '.docx'
            doc.save(fileName)
            self.close()

    def load_from_python(self):
        for section_widget in self.section_widgets:
            section_widget.remove_self()

        dialog = QFileDialog(self)
        options = QFileDialog.options(dialog)
        errors = []
        fileName, _ = QFileDialog.getOpenFileName(
            self,
            "Open Python File",
            "",
            "Python files (*.py);;All Files (*);;",
            options=options
        )

        if fileName:
            try:
                # Load the Python file as a module dynamically
                module_name = os.path.splitext(os.path.basename(fileName))[0]
                spec = importlib.util.spec_from_file_location(
                    module_name, fileName)
                module = importlib.util.module_from_spec(spec)
                spec.loader.exec_module(module)

                # Extract expected variables
                if not hasattr(module, "report_title") or not hasattr(module, "sections"):
                    raise ValueError(
                        "Python file must define 'report_title' and 'sections'")

                self.update_report_title(module.report_title)
                self.update_report_subtitle(module.report_subtitle)
                sections = module.sections.get(self.window_type, [])

                for section in sections:
                    section_name = section.get("section_name")
                    section_widget = self.add_section(section_name)
                    analyses = section.get("analyses", [])

                    for analysis in analyses:
                        analysis_name = settings.ANALYSES[self.window_type].get(
                            analysis.get("analysis"), {}).get("label", "Unknown")
                        channel = analysis.get("channel", "")
                        channel1 = analysis.get("channel1", "")
                        channel2 = analysis.get("channel2", "")

                        invalid_channel = any(
                            c and c not in self.dataMixin.channels for c in [channel, channel1, channel2]
                        )

                        if invalid_channel:
                            errors.append(
                                f"{section_name}/{analysis_name}: Invalid channel name {c}")
                            continue

                        widget = AnalysisWidget(
                            self.main_window, analysis_name, self.window_type, analysis_title=analysis.get("analysis_title"), info_string=analysis.get("info_string"), report_layout=analysis.get("report_layout"))

                        # Assign attributes if they exist
                        for attr in [
                            "channel", "channel1", "channel2",
                            "analysis_range_low", "analysis_range_high",
                            "band_pass_low", "band_pass_high",
                            "show_individual_profiles", "show_min_max", "show_legend",
                            "show_wavelength_labels", "show_unfiltered_data",
                            "machine_speed", "frequency_range_low", "frequency_range_high",
                            "selected_frequencies", "nperseg"
                        ]:
                            if attr in analysis:
                                setattr(widget.controller,
                                        attr, analysis[attr])

                        widget.preview_window.refresh()
                        section_widget.add_analysis(widget)

            except Exception as e:
                traceback.print_exc()
                errors.append(str(e))

            if errors:
                show_json_load_error_msgbox(errors)

    def load_from_json(self):
        for section_widget in self.section_widgets:
            section_widget.remove_self()
        dialog = QFileDialog(self)
        options = QFileDialog.options(dialog)
        errors = []
        fileName, _ = QFileDialog.getOpenFileName(
            self,
            "Open File",
            "",
            "All Files (*);;JSON files (*.json)",
            options=options)
        if fileName:
            with open(fileName, 'r') as file:
                try:
                    data = json.load(file)
                    self.update_report_title(data["report_title"])
                    sections = data["sections"][self.window_type]
                    for section in sections:
                        section_name = section["section_name"]
                        section_widget = self.add_section(section_name)
                        analyses = section["analyses"]
                        for analysis in analyses:
                            analysis_name = settings.ANALYSES[self.window_type][analysis["analysis"]]["label"]
                            channel = analysis.get("channel",  "")
                            channel1 = analysis.get("channel1", "")
                            channel2 = analysis.get("channel2", "")
                            invalid_channel = False
                            for c in [channel, channel1, channel2]:
                                if c and c not in self.dataMixin.channels:
                                    errors.append(
                                        f"{section_name}/{analysis_name}: Invalid channel name {c}")
                                    invalid_channel = True
                                    break  # Break out of the inner channel validation loop
                            if invalid_channel:
                                continue  # Continue to the next analysis in the analyses loop
                            widget = AnalysisWidget(
                                self.main_window, analysis_name, self.window_type)
                            if hasattr(widget.controller, "max_dist"):
                                max_dist = widget.controller.max_dist
                            if hasattr(widget.controller, "max_freq"):
                                max_freq = widget.controller.max_freq
                            if "channel" in analysis:
                                widget.controller.channel = analysis["channel"]
                            if "channel1" in analysis:
                                widget.controller.channel1 = analysis["channel1"]
                            if "channel2" in analysis:
                                widget.controller.channel2 = analysis["channel2"]
                            if "analysis_range_low" in analysis:
                                widget.controller.analysis_range_low = analysis[
                                    "analysis_range_low"] * max_dist
                            if "analysis_range_high" in analysis:
                                widget.controller.analysis_range_high = analysis[
                                    "analysis_range_high"] * max_dist
                            if "band_pass_low" in analysis:
                                widget.controller.band_pass_low = analysis["band_pass_low"]
                            if "band_pass_high" in analysis:
                                widget.controller.band_pass_high = analysis["band_pass_high"]
                            if "show_individual_profiles" in analysis:
                                widget.controller.show_profiles = analysis["show_individual_profiles"]
                            if "show_min_max" in analysis:
                                widget.controller.show_min_max = analysis["show_min_max"]
                            if "show_legend" in analysis:
                                widget.controller.show_legend = analysis["show_legend"]
                            if "show_wavelength_labels" in analysis:
                                widget.controller.show_wavelength = analysis["show_wavelength_labels"]
                            if "show_unfiltered_data" in analysis:
                                widget.controller.show_unfiltered_data = analysis["show_unfiltered_data"]
                            if "machine_speed" in analysis:
                                widget.controller.machine_speed = analysis["machine_speed"]
                            if "frequency_range_low" in analysis:
                                widget.controller.frequency_range_low = analysis[
                                    "frequency_range_low"] * max_freq
                            if "frequency_range_high" in analysis:
                                widget.controller.frequency_range_high = analysis[
                                    "frequency_range_high"] * max_freq
                            if "selected_frequencies" in analysis:
                                widget.controller.selected_freqs = analysis["selected_frequencies"]
                            if "nperseg" in analysis:
                                widget.controller.nperseg = analysis["nperseg"]
                            widget.preview_window.refresh()
                            section_widget.add_analysis(widget)
                except Exception as e:
                    traceback.print_exc()
                    errors.append(str(e))

            if len(errors) > 0:
                show_json_load_error_msgbox(errors)

    def closeEvent(self, event):
        for widget in self.section_widgets:
            widget.remove_self()
        event.accept()


class ReportSectionWidget(QFrame):
    def __init__(self, main_window, section_name="Section", window_type="MD"):
        super().__init__()
        self.main_window = main_window
        self.section_name = section_name
        self.window_type = window_type
        self.analysis_widgets = []

        self.setFrameStyle(QFrame.Shape.StyledPanel | QFrame.Shadow.Plain)
        self.setLineWidth(1)

        self.layout = QVBoxLayout()
        self.layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.layout.setContentsMargins(15, 15, 15, 15)
        self.setLayout(self.layout)

        # Horizontal layout for heading and remove button
        heading_layout = QHBoxLayout()
        self.layout.addLayout(heading_layout)

        self.section_name_label = QLabel("Section name:")
        heading_layout.addWidget(self.section_name_label)

        # Remove button for section
        self.remove_button = QPushButton("Remove section")
        self.remove_button.clicked.connect(self.remove_self)
        heading_layout.addWidget(self.remove_button)

        # Section name input
        self.section_name_input = QLineEdit(self.section_name)
        self.section_name_input.textEdited.connect(self.update_section_name)
        self.section_name_input.setAlignment(Qt.AlignmentFlag.AlignLeft)
        self.layout.addWidget(self.section_name_input)

        self.analyses_label = QLabel("Included analyses:")
        self.layout.addWidget(self.analyses_label)

        # Container for analyses
        self.analyses_container = QVBoxLayout()
        self.analyses_container.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.layout.addLayout(self.analyses_container)

        # ComboBox for analysis items
        self.analysis_combobox = QComboBox()
        self.analyses = settings.ANALYSES[self.window_type]
        labels = [analysis["label"] for analysis in self.analyses.values()]
        self.setup_combobox(self.analysis_combobox, labels)
        self.analysis_combobox.currentIndexChanged.connect(
            self.add_new_analysis)
        self.layout.addWidget(self.analysis_combobox)

    def update_section_name(self, name):
        self.section_name = name

    def setup_combobox(self, combobox, items):
        model = QStandardItemModel()
        add_item = QStandardItem("Add analysis")
        model.appendRow(add_item)
        combobox.insertSeparator(1)

        for item in items:
            model.appendRow(QStandardItem(item))

        combobox.setModel(model)
        combobox.setCurrentIndex(0)

    def add_new_analysis(self):
        if self.analysis_combobox.currentIndex() == 0:
            return

        analysis_name = self.analysis_combobox.currentText()
        if analysis_name:  # Ensure a valid analysis is selected
            analysis_widget = AnalysisWidget(
                self.main_window, analysis_name, self.window_type)
            self.analysis_widgets.append(analysis_widget)
            analysis_widget.destroyed.connect(
                lambda: self.analysis_widgets.remove(analysis_widget))
            self.analyses_container.addWidget(analysis_widget)
            # Optionally clear the combo box to prevent adding the same analysis again
            self.analysis_combobox.setCurrentIndex(0)

    def add_analysis(self, analysis_widget):
        self.analysis_widgets.append(analysis_widget)
        analysis_widget.destroyed.connect(
            lambda: self.analysis_widgets.remove(analysis_widget))
        self.analyses_container.addWidget(analysis_widget)

    def remove_self(self):
        for widget in self.analysis_widgets:
            widget.remove_self()
        self.setParent(None)
        self.deleteLater()


class AnalysisWidget(QWidget):
    def __init__(self, main_window, analysis_name, window_type="MD", analysis_title=None, info_string=None, report_layout=None):
        super().__init__()
        self.analysis_name = analysis_name
        self.window_type = window_type
        self.main_window = main_window
        self.analyses = settings.ANALYSES
        self.analysis_title = analysis_title
        self.info_string = info_string
        self.report_layout = report_layout

        if self.window_type == "CD":
            if analysis_name == self.analyses[self.window_type]["profile"]["label"]:
                window_type = "2d"
                self.controller = CDProfileController(window_type)
                self.preview_window = CDProfileWindow(
                    window_type, self.controller)

            elif analysis_name == self.analyses[self.window_type]["profile_waterfall"]["label"]:
                window_type = "waterfall"
                self.controller = CDProfileController(window_type)
                self.preview_window = CDProfileWindow(
                    window_type, self.controller)

            elif analysis_name == self.analyses[self.window_type]["vca"]["label"]:
                self.controller = VCAController()
                self.preview_window = VCAWindow(self.controller)

        elif self.window_type == "MD":
            if analysis_name == self.analyses[self.window_type]["time_domain"]["label"]:
                self.controller = TimeDomainController()
                self.preview_window = TimeDomainWindow(self.controller)

        if analysis_name == self.analyses[self.window_type]["spectrum"]["label"]:
            self.controller = SpectrumController(self.window_type)
            self.preview_window = SpectrumWindow(
                self.window_type, self.controller)

        elif analysis_name == self.analyses[self.window_type]["spectrogram"]["label"]:
            self.controller = SpectrogramController(self.window_type)
            self.preview_window = SpectrogramWindow(
                self.window_type, self.controller)

        elif analysis_name == self.analyses[self.window_type]["channel_correlation"]["label"]:
            self.controller = ChannelCorrelationController(self.window_type)
            self.preview_window = ChannelCorrelationWindow(
                self.window_type, self.controller)

        elif analysis_name == self.analyses[self.window_type]["correlation_matrix"]["label"]:
            self.controller = CorrelationMatrixController(self.window_type)
            self.preview_window = CorrelationMatrixWindow(
                self.window_type, self.controller)

        elif analysis_name == self.analyses[self.window_type]["formation"]["label"]:
            self.controller = FormationController(self.window_type)
            self.preview_window = FormationWindow(
                self.window_type, self.controller)

        if self.controller:
            self.controller.updated.connect(self.update_analysis_label)
            self.controller.updated.connect(
                lambda: apply_plot_customizations(self))
            apply_plot_customizations(self)

        self.layout = QHBoxLayout()
        self.layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.setLayout(self.layout)

        # Analysis label
        analysis_label_text = self.analysis_title or f"{
            self.analysis_name} {self.get_channel_text()}"

        self.analysis_label = QLabel(analysis_label_text)
        self.analysis_label.setAlignment(Qt.AlignmentFlag.AlignLeft)
        self.layout.addWidget(self.analysis_label)

        # Button layout to align buttons to the right
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        # Preview button
        self.preview_button = QPushButton("Preview")
        self.preview_button.setSizePolicy(
            QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Minimum)
        if self.preview:
            self.preview_button.clicked.connect(self.preview)
        button_layout.addWidget(self.preview_button)

        # Remove button for analysis
        self.remove_button = QPushButton("Remove")
        self.remove_button.setSizePolicy(
            QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Minimum)
        self.remove_button.clicked.connect(self.remove_self)
        button_layout.addWidget(self.remove_button)

        self.layout.addLayout(button_layout)

        # TODO: Fix this, causing flashing windows. But if the window is not viewed, the aspect ratio of the images is also not updated!
        self.preview_window.show()
        self.preview_window.hide()

        self.preview_window.refresh()

    def get_channel_text(self):
        if hasattr(self.controller, "channel"):
            return f"({self.controller.channel})"
        elif hasattr(self.controller, "channel1") and hasattr(self.controller, "channel2"):
            return f"({self.controller.channel1}, {self.controller.channel2})"
        else:
            return ""

    def update_analysis_label(self):
        self.analysis_label.setText(f"{self.analysis_name} {
                                    self.get_channel_text()}")

    def preview(self):
        if not self.preview_window:
            return
        if self.preview_window.isVisible():
            self.preview_window.activateWindow()
        else:
            self.preview_window.show()
            self.main_window.windows.append(self.preview_window)
            self.main_window.updateWindowsList()

    def remove_self(self):
        if self.preview_window:
            self.preview_window.close()
        self.setParent(None)
        self.deleteLater()


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def show_json_load_error_msgbox(errors):
    msgbox = QMessageBox()
    msgbox.setText("Errors occurred while loading report template:")
    msgbox.setInformativeText("\n".join(errors))
    msgbox.exec()
