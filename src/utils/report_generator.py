from abc import ABC, abstractmethod
import datetime
from docx import Document
from docx.shared import Mm, Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pylatex import Document as LatexDocument, Section, Subsection, Figure, NoEscape, Package, Table, Tabular
from pylatex.utils import bold
import os
import shutil
import numpy as np
from .report import get_text_width, set_paragraph_spacing
import settings
import io
import uuid

class ReportGenerator(ABC):
    def __init__(self, report_data):
        self.report_data = report_data
        self.report_title = report_data.get('title', '')
        self.report_subtitle = report_data.get('subtitle', '')
        self.additional_info = report_data.get('additional_info', '')
        self.header_image_path = report_data.get('header_image_path', '')
        self.sample_image_path = report_data.get('sample_image_path', '')
        self.sections = report_data.get('sections', [])
        self.window_type = report_data.get('window_type', 'MD')

    @abstractmethod
    def generate(self, output_path):
        """Generate the report and save it to the specified path"""
        pass

class WordReportGenerator(ReportGenerator):
    def generate(self, output_path):
        doc = Document()
        margin = Cm(2)
        for section in doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        style = doc.styles['Normal']
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        # Explicitly define font in Word
        r = style._element
        r_rPr = r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "DejaVu Sans")
        rFonts.set(qn("w:hAnsi"), "DejaVu Sans")
        rFonts.set(qn("w:eastAsia"), "DejaVu Sans")
        rFonts.set(qn("w:cs"), "DejaVu Sans")
        r_rPr.append(rFonts)

        # Header section
        header = doc.sections[0].header
        total_width_mm = get_text_width(doc)
        header_table = header.add_table(rows=1, cols=3, width=Mm(total_width_mm))
        header_table.autofit = False

        # Distribute widths
        left_col_width = Mm(total_width_mm * 0.25)
        right_col_width = Mm(total_width_mm * 0.25)
        middle_col_width = Mm(total_width_mm * 0.50)

        header_table.columns[0].width = left_col_width
        header_table.columns[1].width = middle_col_width
        header_table.columns[2].width = right_col_width

        cell1 = header_table.cell(0, 0)
        cell2 = header_table.cell(0, 1)
        cell3 = header_table.cell(0, 2)

        # Add header image
        if self.header_image_path:
            run = cell1.paragraphs[0].add_run()
            run.add_picture(self.header_image_path, width=Mm(left_col_width.mm * 0.9))

        # Add title and info
        p = cell2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.clear()
        
        run_title = p.add_run(self.report_title)
        run_title.bold = True

        run_subtitle = p.add_run(f"\n{self.report_subtitle}\n")
        run_subtitle.bold = False

        run_additional_info = p.add_run(f"\n{self.additional_info}")
        run_additional_info.bold = False

        run_subtitle = p.add_run(f"\nGenerated {str(datetime.datetime.now())}")

        # Add sample image
        if self.sample_image_path:
            cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = cell3.paragraphs[0].add_run()
            run.add_picture(self.sample_image_path, width=Mm(right_col_width.mm * 0.9))

        # Add sections
        for section in self.sections:
            self._add_section_to_word(doc, section)

        # Save document
        doc.save(output_path)

    def _add_section_to_word(self, doc, section):
        paragraph = doc.add_heading(section.section_name)
        run = paragraph.runs[0]
        run.font.color.rgb = None
        run.font.size = Pt(12)
        set_paragraph_spacing(paragraph, 0, 6)

        for analysis in section.analysis_widgets:
            if settings.REPORT_ENABLE_ANALYSIS_TITLE:
                analysis_title = analysis.analysis_title or f"{analysis.analysis_name} {analysis.get_channel_text()}"
                paragraph = doc.add_heading(analysis_title, 2)

            if analysis.info_string:
                paragraph = doc.add_paragraph()
                run_info = paragraph.add_run(analysis.info_string)
                run_info.italic = True
                run_info.font.size = Pt(8)

            set_paragraph_spacing(paragraph)
            self._add_analysis_to_word(doc, analysis)

        if doc.paragraphs[-1].text.strip():
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def _add_analysis_to_word(self, doc, analysis):
        layout_mode = analysis.report_layout or "stats-right"
        total_width_mm = get_text_width(doc)
        img_col_width = Mm(total_width_mm * (3/5))
        stats_col_width = Mm(total_width_mm * (2/5))

        if layout_mode == "stats-right":
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            
            # Image column
            col1 = table.columns[0]
            col1.width = img_col_width
            cell1 = table.cell(0, 0)
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            run = cell1.paragraphs[0].add_run()
            run.add_picture(analysis.controller.getPlotImage(), width=Cm(col1.width.cm - 0.5))

            # Stats column
            col2 = table.columns[1]
            col2.width = stats_col_width
            cell2 = table.cell(0, 1)
            cell2.width = stats_col_width
            cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            self._add_stats_table(cell2, analysis)

        elif layout_mode in ["stats-below", "stats-above"]:
            table = doc.add_table(rows=2, cols=1)
            table.autofit = True

            if layout_mode == "stats-above":
                stats_cell = table.cell(0, 0)
                img_cell = table.cell(1, 0)
            else:
                img_cell = table.cell(0, 0)
                stats_cell = table.cell(1, 0)

            run = img_cell.paragraphs[0].add_run()
            run.add_picture(analysis.controller.getPlotImage(), 
                          width=Mm(analysis.image_width_mm or (total_width_mm - 5)))
            img_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            stats_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            self._add_stats_table(stats_cell, analysis)

    def _add_stats_table(self, cell, analysis):
        data = analysis.controller.getStatsTableData()
        if not data:
            return

        shape = np.shape(data)
        rows, cols = shape if len(shape) == 2 else (shape[0], 1) if len(shape) == 1 else (1, 1)

        stats_table = cell.add_table(rows, cols)
        cell.paragraphs[0].clear()
        stats_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row_idx, row_data in enumerate(data):
            row = stats_table.rows[row_idx]
            for col_idx, cell_data in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.paragraphs[0].text = cell_data
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell.paragraphs[0].style.font.size = Pt(8)
                cell.paragraphs[0].style.font.name = "Nimbus Mono PS"
                cell.width = Mm(30)

class LatexReportGenerator(ReportGenerator):
    def generate(self, output_path):
        # Create output directory if it doesn't exist
        output_dir = os.path.dirname(output_path) or '.'
        images_dir = os.path.join(output_dir, 'images')
        os.makedirs(images_dir, exist_ok=True)

        # Create document
        geometry_options = {
            "margin": "2cm",
            "includeheadfoot": True
        }
        doc = LatexDocument(documentclass='article', geometry_options=geometry_options)
        
        # Add required packages
        doc.packages.append(Package('graphicx'))
        doc.packages.append(Package('booktabs'))
        doc.packages.append(Package('float'))
        doc.packages.append(Package('fancyhdr'))
        doc.packages.append(Package('array'))
        doc.packages.append(Package('caption'))
        doc.packages.append(Package('siunitx'))  # For better table number formatting
        
        # Configure page style
        doc.preamble.append(NoEscape(r'\pagestyle{fancy}'))
        doc.preamble.append(NoEscape(r'\fancyhead{}'))  # Clear all header fields
        doc.preamble.append(NoEscape(r'\graphicspath{{./images/}}'))  # Set graphics path
        
        # Add title
        doc.preamble.append(NoEscape(r'\title{' + self._escape_latex(self.report_title) + r'}'))
        doc.preamble.append(NoEscape(r'\author{' + self._escape_latex(self.report_subtitle) + r'}'))
        doc.preamble.append(NoEscape(r'\date{\today}'))
        
        doc.append(NoEscape(r'\maketitle'))
        
        # Add header images if they exist
        if self.header_image_path or self.sample_image_path:
            with doc.create(Figure(position='H')) as fig:
                if self.header_image_path:
                    header_img = self._copy_and_convert_image(self.header_image_path, images_dir, 'header')
                    if header_img:
                        fig.add_image('header', width='0.4\\textwidth')
                if self.sample_image_path:
                    sample_img = self._copy_and_convert_image(self.sample_image_path, images_dir, 'sample')
                    if sample_img:
                        fig.add_image('sample', width='0.4\\textwidth')
                fig.add_caption('Header Images')
        
        # Add additional info
        if self.additional_info:
            doc.append(NoEscape(r'\textit{' + self._escape_latex(self.additional_info) + r'}\par\vspace{1em}'))
        
        # Add sections
        for section in self.sections:
            self._add_section_to_latex(doc, section, images_dir)
        
        # Generate PDF
        if settings.REPORT_GENERATE_PDF:
            try:
                doc.generate_pdf(os.path.splitext(output_path)[0], clean_tex=False)
            except Exception as e:
                # If PDF generation fails, at least save the .tex file
                doc.generate_tex(os.path.splitext(output_path)[0])
                raise Exception(f"Failed to generate PDF. LaTeX file saved. Error: {str(e)}")
        else:
            doc.generate_tex(os.path.splitext(output_path)[0])

    def _add_section_to_latex(self, doc, section, images_dir):
        with doc.create(Section(self._escape_latex(section.section_name))):
            for analysis in section.analysis_widgets:
                if settings.REPORT_ENABLE_ANALYSIS_TITLE:
                    analysis_title = analysis.analysis_title or f"{analysis.analysis_name} {analysis.get_channel_text()}"
                    with doc.create(Subsection(self._escape_latex(analysis_title))):
                        self._add_analysis_to_latex(doc, analysis, images_dir)
                else:
                    self._add_analysis_to_latex(doc, analysis, images_dir)

    def _add_analysis_to_latex(self, doc, analysis, images_dir):
        # Add info string if it exists
        if analysis.info_string:
            doc.append(NoEscape(r'\textit{' + self._escape_latex(analysis.info_string) + r'}\par\vspace{0.5em}'))

        # Generate unique identifier for this analysis
        analysis_id = str(uuid.uuid4())[:8]

        # Add plot image
        plot_filename = f"plot_{analysis.analysis_name.lower().replace(' ', '_')}_{analysis_id}"
        with doc.create(Figure(position='H')) as fig:
            # Save plot image
            plot_img = self._save_plot_image(analysis.controller.getPlotImage(format="pdf"), images_dir, plot_filename, format="pdf")
            if plot_img:
                fig.add_image(plot_filename, width=NoEscape('0.8\\textwidth'))
                fig.add_caption(self._escape_latex(analysis.analysis_name))

        # Add stats table
        data = analysis.controller.getStatsTableData()
        if data:
            # Get number of columns
            shape = np.shape(data)
            num_cols = shape[1] if len(shape) > 1 else 1
            
            # Create table with booktabs style
            with doc.create(Table(position='H')) as table:
                # Create tabular environment with column specifications
                col_spec = '|' + '|'.join(['l' if i == 0 else 'r' for i in range(num_cols)]) + '|'
                tabular = Tabular(col_spec)
                table.append(NoEscape(r'\centering'))
                table.append(tabular)
                
                # Add the data
                tabular.add_hline()
                for row_data in data:
                    # Convert all data to strings and escape special characters
                    escaped_row = []
                    for i, cell in enumerate(row_data):
                        cell_str = str(cell).strip()
                        # Try to format numbers with siunitx if it's a number
                        try:
                            float(cell_str)  # Check if it's a number
                            escaped_row.append(NoEscape(rf'\num{{{cell_str}}}'))
                        except ValueError:
                            escaped_row.append(self._escape_latex(cell_str))
                    
                    tabular.add_row(escaped_row)
                    tabular.add_hline()
            
            # Add vertical space after the table using \vspace
            doc.append(NoEscape(r'\vspace{1em}'))

    def _escape_latex(self, text):
        """Escape special LaTeX characters"""
        if not isinstance(text, str):
            text = str(text)
        
        # First handle backslashes and braces
        text = text.replace('\\', r'\textbackslash{}')
        text = text.replace('{', r'\{').replace('}', r'\}')
        
        # Then handle other special characters
        latex_special_chars = {
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '~': r'\textasciitilde{}',
            '^': r'\^{}',
            '<': r'\textless{}',
            '>': r'\textgreater{}',
        }
        return ''.join(latex_special_chars.get(c, c) for c in text)

    def _copy_and_convert_image(self, src_path, dest_dir, name):
        """Copy image to destination directory and return the new filename"""
        if not src_path or not os.path.exists(src_path):
            return None
            
        # Get file extension and create new filename
        _, ext = os.path.splitext(src_path)
        new_filename = f"{name}{ext}"
        dest_path = os.path.join(dest_dir, new_filename)
        
        # Copy file
        shutil.copy2(src_path, dest_path)
        return name  # Return basename without extension for LaTeX

    def _save_plot_image(self, buffer, dest_dir, name, format="png"):
        """Save plot from BytesIO buffer to destination directory"""
        if not buffer or not isinstance(buffer, io.BytesIO):
            return None
            
        # Save as PNG
        dest_path = os.path.join(dest_dir, f"{name}.{format}")
        
        # Write buffer contents to file
        with open(dest_path, 'wb') as f:
            f.write(buffer.getvalue())
            
        return name  # Return basename without extension for LaTeX

def create_report_generator(report_type, report_data):
    """Factory function to create appropriate report generator"""
    generators = {
        'word': WordReportGenerator,
        'latex': LatexReportGenerator
    }
    
    generator_class = generators.get(report_type.lower())
    if not generator_class:
        raise ValueError(f"Unsupported report type: {report_type}")
    
    return generator_class(report_data)
