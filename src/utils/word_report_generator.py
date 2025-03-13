from docx import Document
from docx.shared import Mm, Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime
import numpy as np
from .report_generator import ReportGenerator
import settings


class WordReportGenerator(ReportGenerator):
    def generate(self, output_path):
        doc = Document()
        margin = Cm(2)
        for section in doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        style = doc.styles['Normal']
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        # Explicitly define font in Word
        r = style._element
        r_rPr = r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "DejaVu Sans")
        rFonts.set(qn("w:hAnsi"), "DejaVu Sans")
        rFonts.set(qn("w:eastAsia"), "DejaVu Sans")
        rFonts.set(qn("w:cs"), "DejaVu Sans")
        r_rPr.append(rFonts)

        # Header section
        header = doc.sections[0].header
        total_width_mm = get_text_width(doc)
        header_table = header.add_table(
            rows=1, cols=3, width=Mm(total_width_mm))
        header_table.autofit = True

        # Distribute widths
        left_col_width = Mm(total_width_mm * 0.25)
        right_col_width = Mm(total_width_mm * 0.25)
        middle_col_width = Mm(total_width_mm * 0.50)

        header_table.columns[0].width = left_col_width
        header_table.columns[1].width = middle_col_width
        header_table.columns[2].width = right_col_width

        cell1 = header_table.cell(0, 0)
        cell2 = header_table.cell(0, 1)
        cell3 = header_table.cell(0, 2)

        # Add header image
        if self.header_image_path:
            run = cell1.paragraphs[0].add_run()
            run.add_picture(self.header_image_path,
                            width=Mm(left_col_width.mm * 0.9))

        # Add title and info
        p = cell2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.clear()

        run_title = p.add_run(self.report_title)
        run_title.bold = True

        run_subtitle = p.add_run(f"\n{self.report_subtitle}\n")
        run_subtitle.bold = False

        run_additional_info = p.add_run(f"\n{self.additional_info}")
        run_additional_info.bold = False

        # run_subtitle = p.add_run(f"\nGenerated {str(datetime.datetime.now())}")

        # Add sample image
        if self.sample_image_path:
            cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = cell3.paragraphs[0].add_run()
            run.add_picture(self.sample_image_path,
                            width=Mm(right_col_width.mm * 0.9))

        # Add sections
        for section in self.sections:
            self._add_section_to_word(doc, section)

        # Save document
        doc.save(output_path)

    def _add_section_to_word(self, doc, section):
        paragraph = doc.add_heading(section.section_name)
        run = paragraph.runs[0]
        run.font.color.rgb = None
        run.font.size = Pt(12)
        set_paragraph_spacing(paragraph, 0, 6)

        for analysis in section.analysis_widgets:
            if settings.REPORT_ENABLE_ANALYSIS_TITLE:
                analysis_title = analysis.analysis_title or f"{analysis.analysis_name} {analysis.get_channel_text()}"
                paragraph = doc.add_heading(analysis_title, 2)

            if analysis.info_string:
                paragraph = doc.add_paragraph()
                run_info = paragraph.add_run(analysis.info_string)
                run_info.italic = True
                run_info.font.size = Pt(8)

            set_paragraph_spacing(paragraph)
            self._add_analysis_to_word(doc, analysis)

        if doc.paragraphs[-1].text.strip():
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def _add_analysis_to_word(self, doc, analysis):
        layout_mode = analysis.report_layout or "stats-right"
        total_width_mm = get_text_width(doc)
        img_col_width = Mm(total_width_mm * (3/5))
        stats_col_width = Mm(total_width_mm * (2/5))

        if layout_mode == "stats-right":
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True

            # Image column
            col1 = table.columns[0]
            col1.width = img_col_width
            cell1 = table.cell(0, 0)
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            run = cell1.paragraphs[0].add_run()
            run.add_picture(analysis.controller.getPlotImage(),
                            width=Cm(col1.width.cm - 0.5))

            # Stats column
            col2 = table.columns[1]
            col2.width = stats_col_width
            cell2 = table.cell(0, 1)
            cell2.width = stats_col_width
            cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            # Clear the default paragraph to prevent extra spacing
            cell2.paragraphs[0].clear()
            self._add_stats_table(cell2, analysis)

        elif layout_mode in ["stats-below", "stats-above"]:
            table = doc.add_table(rows=2, cols=1)
            table.autofit = True

            if layout_mode == "stats-above":
                stats_cell = table.cell(0, 0)
                img_cell = table.cell(1, 0)
            else:
                img_cell = table.cell(0, 0)
                stats_cell = table.cell(1, 0)

            run = img_cell.paragraphs[0].add_run()
            run.add_picture(analysis.controller.getPlotImage(),
                            width=Mm(analysis.image_width_mm or (total_width_mm - 5)))
            img_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            stats_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            self._add_stats_table(stats_cell, analysis)

    def _add_stats_table(self, cell, analysis):
        data = analysis.controller.getStatsTableData()
        if not data:
            return

        shape = np.shape(data)
        rows, cols = shape if len(shape) == 2 else (
            shape[0], 1) if len(shape) == 1 else (1, 1)

        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)
    


        stats_table = cell.add_table(rows, cols)
        cell.paragraphs[0].clear()
        stats_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row_idx, row_data in enumerate(data):
            row = stats_table.rows[row_idx]
            for col_idx, cell_data in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.paragraphs[0].text = cell_data
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell.paragraphs[0].style.font.size = Pt(8)
                cell.paragraphs[0].style.font.name = "Nimbus Mono PS"
                cell.width = Mm(30)


def set_paragraph_spacing(paragraph, space_before=0, space_after=0, line_spacing=1):
    """
    Set the spacing for a paragraph.

    Parameters:
    paragraph (docx.text.paragraph.Paragraph): The paragraph to format.
    space_before (float): Space before the paragraph in points.
    space_after (float): Space after the paragraph in points.
    line_spacing (float): Line spacing, where 1 is single, 2 is double, etc.
    """
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)
    paragraph_format.line_spacing = line_spacing


def get_text_width(document):
    """
    Returns the text width in mm.
    """
    section = document.sections[0]
    return (section.page_width - section.left_margin - section.right_margin) / 36000
